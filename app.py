from docx import Document
from pyphen import Pyphen

word = "Dictionary"

dic = Pyphen(lang="pl_PL")

word_pyphened = dic.inserted(word)

doc = Document()
p = doc.add_paragraph()

rnr = p.add_run(word_pyphened)
rnr.bold = True

doc.save("test.docx")